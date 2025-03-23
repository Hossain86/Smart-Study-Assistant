import re
import os
import io
import tempfile
from flask import Flask, request, jsonify, send_file
import fitz  # PyMuPDF
from dotenv import load_dotenv
from flask_cors import CORS
from docx import Document
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
from assignment_generator import assignment_gen, markdown_to_plain_text
from coverpage_generator import generate_coverpage
from mcq_generator import generateMCQ
from narrative_generator import generateOpenEnded

load_dotenv()

app = Flask(__name__)
CORS(app) 

def extractText(pdf_path):
    """Extract full text from a given PDF file."""
    doc = fitz.open(pdf_path)
    text = "\n".join([page.get_text("text") for page in doc])
    return text.strip()

@app.route("/extract-text", methods=["POST"])
def extract_text_endpoint():
    """API endpoint to extract and return the text from an uploaded PDF."""
    
    if 'pdf' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    pdf_file = request.files['pdf']

    if pdf_file.filename == "":
        return jsonify({"error": "Empty filename"}), 400

    pdf_path = "temp.pdf"
    pdf_file.save(pdf_path)
    extracted_text = extractText(pdf_path)

    os.remove(pdf_path)

    if not extracted_text.strip():
        return jsonify({"error": "No text found in PDF"}), 400

    return jsonify({"extracted_text": extracted_text})

@app.route('/')
def home():
    return "<h1>Welcome to the PDF Question Generator API!</h1>"

@app.route("/generate-questions", methods=["POST"])
def generate_questions():
    """API endpoint to generate MCQs or OpenEnded Questions from an uploaded PDF."""
    
    if 'pdf' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    pdf_file = request.files['pdf']
    question_type = request.form.get("question_type", "mcq")

    if pdf_file.filename == "":
        return jsonify({"error": "Empty filename"}), 400

    pdf_path = "temp.pdf"
    pdf_file.save(pdf_path)
    extracted_text = extractText(pdf_path)

    if not extracted_text.strip():
        return jsonify({"error": "No text found in PDF"}), 400

    print(f"Received request for {question_type}")  # Debug log

    result = generateMCQ(extracted_text) if question_type == "mcq" else generateOpenEnded(extracted_text)

    os.remove(pdf_path)

    return jsonify(result)


@app.route("/generate-lab-report", methods=["POST"])
def generate_Lab_Report():
    """API endpoint to generate a lab report from a given topic."""
    if not request.is_json:
        return jsonify({"error": "Invalid request. JSON data expected."}), 400
    
    data = request.get_json()
    topic = data.get("topic")
    if not topic:
        return jsonify({"error": "Missing topic"}), 400
    print(f"Received request for Lab Report on: {topic}")
    lab_report = assignment_gen(topic)
    if "error" in lab_report:
        return jsonify({"error": lab_report["error"]}), 500

    return jsonify(lab_report)

@app.route("/lab-report-docx", methods=["POST"])
def lab_report_docx():
    """API endpoint to generate a DOCX file for the lab report from a given topic."""
    if not request.is_json:
        return jsonify({"error": "Invalid request. JSON data expected."}), 400

    data = request.get_json()
    topic = data.get("topic")
    if not topic:
        return jsonify({"error": "Missing topic"}), 400

    print(f"Generating DOCX Lab Report for: {topic}")
    # Generate lab report using assignment_gen function
    lab_report_md = assignment_gen(topic)
    if isinstance(lab_report_md, dict) and "error" in lab_report_md:
        return jsonify({"error": lab_report_md["error"]}), 500

    # Convert Markdown to plain text
    plain_text = markdown_to_plain_text(lab_report_md)

    # Create a DOCX document from the plain text
    document = Document()
    for line in plain_text.split('\n'):
        document.add_paragraph(line)

    # Save the document into an in-memory bytes buffer
    file_buffer = io.BytesIO()
    document.save(file_buffer)
    file_buffer.seek(0)

    # Send the DOCX file as a downloadable attachment
    return send_file(
        file_buffer,
        as_attachment=True,
        download_name="lab_report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/final-docx", methods=["POST"])
def final_docx():
    """
    API endpoint to generate a final DOCX that includes both a coverpage and a lab report.
    Expects JSON with:
      - topic: string
      - cover_data: object containing coverpage fields (optional; defaults are provided)
    """
    if not request.is_json:
        return jsonify({"error": "Invalid request. JSON data expected."}), 400

    data = request.get_json()
    topic = data.get("topic")
    if not topic:
        return jsonify({"error": "Missing topic"}), 400

    # Extract coverpage data from the request; use defaults if not provided.
    cover_data = data.get("cover_data", {})
    defaults = {
        "department": "Computer Science and Engineering",
        "course_code": "CSE-1102",
        "course_name": "Introduction to Programming Sessional",
        "assignment_name": "Lab Report Assignment",
        "date_of_submission": "01/01/2025",
        "submitted_by_name": "Student Name",
        "submitted_by_roll": "2103000",
        "submitted_by_section": "A",
        "submitted_by_series": "21",
        "submitted_to": "Dr. Teacher"
    }
    for key, value in defaults.items():
        if key not in cover_data or not cover_data[key]:
            cover_data[key] = value

    # 1. Generate Lab Report DOCX from markdown
    lab_report_md = assignment_gen(topic)
    if isinstance(lab_report_md, dict) and "error" in lab_report_md:
        return jsonify({"error": lab_report_md["error"]}), 500

    plain_text = markdown_to_plain_text(lab_report_md)
    lab_report_doc = Document()
    for line in plain_text.split('\n'):
        lab_report_doc.add_paragraph(line)

    # 2. Generate Coverpage DOCX using the coverpage template
    coverpage_doc_tpl = generate_coverpage(cover_data, template_path="coverpage_template.docx")
    # Save the rendered coverpage to a temporary file.
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_cover:
        coverpage_doc_tpl.save(temp_cover.name)
    
    # 3. Save lab report document to a temporary file.
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_lab:
        lab_report_doc.save(temp_lab.name)

    # 4. Merge the two documents using docxcompose (preserving each document's formatting)
    coverpage_doc = Document(temp_cover.name)
    lab_report_doc = Document(temp_lab.name)
    composer = Composer(coverpage_doc)
    composer.append(lab_report_doc)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_final:
        composer.save(temp_final.name)
        temp_final.seek(0)
        final_data = temp_final.read()

    # Return the final merged document as a downloadable file.
    return send_file(
        io.BytesIO(final_data),
        as_attachment=True,
        download_name="final_lab_report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    
    
@app.route('/submit', methods=['POST'])
def submit_data():
    data = request.get_json()
    print(data)
    return jsonify({"message": "Data received!", "data": data})

if __name__ == "__main__":
    app.run(debug=True)
